import os
import datetime
from docx import Document

# Directory where VS Code files are stored
WORKSPACE_DIR = "path_to_student_workspace"
# Path to the docx template
TEMPLATE_PATH = "path_to_docx_template/progress_report_template.docx"
# Directory to save the generated reports
REPORT_DIR = "path_to_save_reports"

# Function to analyze VS Code workspace
def analyze_workspace(workspace_dir):
    student_data = {
        "total_files": 0,
        "files_list": [],
        "last_modified": None,
    }

    for root, _, files in os.walk(workspace_dir):
        for file in files:
            student_data["total_files"] += 1
            file_path = os.path.join(root, file)
            student_data["files_list"].append(file_path)
            modified_time = os.path.getmtime(file_path)
            if not student_data["last_modified"] or modified_time > student_data["last_modified"]:
                student_data["last_modified"] = modified_time

    if student_data["last_modified"]:
        student_data["last_modified"] = datetime.datetime.fromtimestamp(student_data["last_modified"]).strftime("%Y-%m-%d %H:%M:%S")
    else:
        student_data["last_modified"] = "No recent activity found"
    
    return student_data

# Function to generate progress report using a docx template
def create_progress_report(student_name, student_data, template_path, report_dir):
    # Load the template
    doc = Document(template_path)

    # Replace placeholders in the template
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("{student_name}", student_name)
        paragraph.text = paragraph.text.replace("{date}", datetime.datetime.now().strftime("%Y-%m-%d"))
        paragraph.text = paragraph.text.replace("{total_files}", str(student_data["total_files"]))
        paragraph.text = paragraph.text.replace("{last_activity}", student_data["last_modified"])

    # Replace files list placeholder
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{files_list}" in cell.text:
                    cell.text = "\n".join(student_data["files_list"][:10])  # Show top 10 files

    # Save the report
    os.makedirs(report_dir, exist_ok=True)
    report_filename = f"{student_name}_Progress_Report_{datetime.datetime.now().strftime('%Y-%m-%d')}.docx"
    report_path = os.path.join(report_dir, report_filename)
    doc.save(report_path)

    return report_path

# Main function to execute automation
def generate_report():
    student_name = "John Doe"  # Replace with dynamic input if needed
    student_data = analyze_workspace(WORKSPACE_DIR)

    if not student_data["files_list"]:
        print(f"No activity found in the workspace: {WORKSPACE_DIR}")
        return

    report_path = create_progress_report(student_name, student_data, TEMPLATE_PATH, REPORT_DIR)
    print(f"Progress report generated: {report_path}")

# Schedule automation (use a scheduler like cron or Windows Task Scheduler to run this script on the 3rd of every month)
if __name__ == "__main__":
    generate_report()
